from __future__ import annotations

import argparse
import os
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

SCRIPT_DIR = Path(__file__).resolve().parent
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

from export_to_google_docs import (
    DEFAULT_GOOGLE_DRIVE_FOLDER_ID,
    build_credentials,
    load_environment,
    upload_google_doc,
    validate_args as validate_export_args,
)


BRAND_NAVY = RGBColor(0x17, 0x16, 0x2C)
BRAND_SLATE = RGBColor(0x24, 0x29, 0x38)
BRAND_MUTED = RGBColor(0x6B, 0x73, 0x88)
BRAND_SOFT = RGBColor(0xF4, 0xF5, 0xF9)

CALLOUT_LABELS = {
    "Huidige status vs. doel",
    "Wat werkt goed?",
    "Wat liep stroef?",
    "Concrete conclusie",
    "Concrete actiepunten",
}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Convert a Markdown audit to a branded DOCX and publish it as a Google Doc."
    )
    parser.add_argument("input_markdown", type=Path, help="Path to the Markdown source audit")
    parser.add_argument(
        "--output-docx",
        type=Path,
        help="Optional path for the intermediate DOCX export. Defaults to output/docx/<input-name>.docx",
    )
    parser.add_argument(
        "--name",
        help="Optional Google Doc filename. Defaults to the Markdown filename without extension.",
    )
    parser.add_argument(
        "--folder-id",
        help="Optional Google Drive folder id. Falls back to the default repo folder or env.",
    )
    parser.add_argument(
        "--service-account-file",
        help="Optional Google service account JSON file.",
    )
    parser.add_argument(
        "--client-secrets-file",
        help="Optional OAuth desktop client credentials JSON file.",
    )
    parser.add_argument(
        "--token-file",
        help="Optional OAuth token file path.",
    )
    parser.add_argument(
        "--oauth-mode",
        choices=("local-server", "console"),
        help="How to complete the first OAuth login.",
    )
    return parser.parse_args()


def build_output_docx_path(input_markdown: Path, output_docx: Path | None) -> Path:
    if output_docx is not None:
        output_docx.parent.mkdir(parents=True, exist_ok=True)
        return output_docx

    default_path = Path("output/docx") / f"{input_markdown.stem}.docx"
    default_path.parent.mkdir(parents=True, exist_ok=True)
    return default_path


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_border(paragraph, color: str = "D6DAE4") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)

    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)

    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)


def configure_styles(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.1)
    section.bottom_margin = Cm(1.9)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    for style_name, font_name, size, color, bold in [
        ("Normal", "Poppins", Pt(10.8), BRAND_SLATE, False),
        ("Title", "Poppins Light", Pt(26), BRAND_NAVY, False),
        ("Subtitle", "Poppins", Pt(11.5), BRAND_MUTED, False),
        ("Heading 1", "Poppins SemiBold", Pt(18), BRAND_NAVY, True),
        ("Heading 2", "Poppins SemiBold", Pt(13), BRAND_NAVY, True),
        ("Heading 3", "Poppins SemiBold", Pt(11.5), BRAND_NAVY, True),
        ("List Bullet", "Poppins", Pt(10.8), BRAND_SLATE, False),
        ("List Number", "Poppins", Pt(10.8), BRAND_SLATE, False),
    ]:
        style = document.styles[style_name]
        style.font.name = font_name
        style.font.size = size
        style.font.color.rgb = color
        style.font.bold = bold

    normal = document.styles["Normal"].paragraph_format
    normal.space_after = Pt(7)
    normal.line_spacing = 1.18

    for style_name, before, after in [
        ("Heading 1", 16, 8),
        ("Heading 2", 12, 6),
        ("Heading 3", 8, 4),
    ]:
        paragraph_format = document.styles[style_name].paragraph_format
        paragraph_format.space_before = Pt(before)
        paragraph_format.space_after = Pt(after)
        paragraph_format.keep_with_next = True


def add_inline_runs(paragraph, text: str) -> None:
    for part in re.split(r"(\*\*.*?\*\*|`[^`]+`)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = "Poppins SemiBold"
            run.font.color.rgb = BRAND_NAVY
        elif part.startswith("`") and part.endswith("`") and len(part) >= 3:
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9.8)
        else:
            run = paragraph.add_run(part)
            run.font.name = "Poppins"
            run.font.color.rgb = BRAND_SLATE


def add_body_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="Normal")
    add_inline_runs(paragraph, text)


def add_special_label(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="Heading 2")
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.font.name = "Poppins SemiBold"
    run.font.size = Pt(12.5)
    run.font.color.rgb = BRAND_NAVY


def add_vs_label(document: Document) -> None:
    paragraph = document.add_paragraph(style="Subtitle")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run("VS")
    run.font.name = "Poppins SemiBold"
    run.font.size = Pt(11)
    run.font.color.rgb = BRAND_MUTED


def add_separator(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(14)
    run = paragraph.add_run("")
    run.font.name = "Poppins"
    set_paragraph_border(paragraph)


def prettify_company_name(name: str) -> str:
    letters_only = re.sub(r"[^A-Za-z]+", "", name)
    if letters_only and letters_only.isupper():
        return name.title()
    return name


def read_related_audit_metadata(input_markdown: Path) -> dict[str, str]:
    match = re.match(r"^(?P<prefix>.+)-seo-geo-klantversie$", input_markdown.stem)
    if not match:
        return {}

    prefix = match.group("prefix")
    candidates = sorted(input_markdown.parent.glob(f"{prefix}-seo-audit-*.md"))
    if not candidates:
        return {}

    metadata: dict[str, str] = {}
    for line in candidates[-1].read_text(encoding="utf-8").splitlines():
        stripped = line.strip()
        match = re.match(r"^\*\*(.+?):\*\*\s*(.+)$", stripped)
        if not match:
            continue
        metadata[match.group(1).strip().lower()] = match.group(2).strip()
    return metadata


def extract_cover_metadata(lines: list[str], input_markdown: Path) -> dict[str, str]:
    fallback_name = input_markdown.stem
    title = ""
    company = ""
    website = ""
    audit_date = ""
    regions = ""

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("# ") and not title:
            title = stripped[2:].strip()
            company = title.split("|", 1)[0].strip()
            company = re.sub(r"\b20\d{2}\b", "", company).strip(" -")
            continue

        match = re.match(r"^\*\*(.+?):\*\*\s*(.+)$", stripped)
        if not match:
            continue
        key = match.group(1).strip().lower()
        value = match.group(2).strip()
        if key == "bedrijf":
            company = value
        elif key == "website":
            website = value
        elif key == "auditdatum":
            audit_date = value
        elif key == "doelregio's":
            regions = value

    related_audit_metadata = read_related_audit_metadata(input_markdown)
    company = company or related_audit_metadata.get("bedrijf", "")
    website = website or related_audit_metadata.get("website", "")
    audit_date = audit_date or related_audit_metadata.get("auditdatum", "")
    regions = regions or related_audit_metadata.get("doelregio's", "")

    if not title:
        title = fallback_name.replace("-", " ").title()
    if not company:
        company = fallback_name.split("-seo-", 1)[0].replace("-", " ").title()
    company = prettify_company_name(company)

    return {
        "title": title,
        "company": company,
        "website": website,
        "audit_date": audit_date,
        "regions": regions,
    }


def build_cover(document: Document, metadata: dict[str, str]) -> None:
    section = document.sections[0]
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = section.page_width - section.left_margin - section.right_margin

    row = table.rows[0]
    row.height = Cm(22.7)
    cell = row.cells[0]
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, "17162C")

    first = cell.paragraphs[0]
    first.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = first.add_run("SEO & GEO KLANTVERSIE 2026")
    run.font.name = "Poppins Light"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xD6, 0xDA, 0xE4)

    brand = cell.add_paragraph()
    brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = brand.add_run(metadata["company"])
    run.font.name = "Poppins Light"
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    title = cell.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SEO & GEO Analyse")
    run.font.name = "Poppins SemiBold"
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    subtitle = cell.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Klantversie")
    run.font.name = "Poppins"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xCB, 0xD2, 0xE2)

    details = [value for value in [metadata["website"], metadata["audit_date"], metadata["regions"]] if value]
    for value in details:
        detail = cell.add_paragraph()
        detail.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = detail.add_run(value)
        run.font.name = "Poppins"
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0xCB, 0xD2, 0xE2)


def build_summary_callout(document: Document, metadata: dict[str, str]) -> None:
    title = metadata["title"]
    company = metadata["company"]
    message = f"{company} krijgt hier een klantversie van de SEO & GEO analyse met focus op prioriteiten, groeikansen en concrete vervolgstappen."
    if "SEO & GEO" not in title.upper():
        message = f"{company} krijgt hier een klantversie van {title.lower()} met focus op prioriteiten, groeikansen en concrete vervolgstappen."

    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F5F9")

    heading = cell.paragraphs[0]
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = heading.add_run("Belangrijkste boodschap")
    run.font.name = "Poppins SemiBold"
    run.font.size = Pt(10.5)
    run.font.color.rgb = BRAND_NAVY

    body = cell.add_paragraph()
    body.paragraph_format.space_before = Pt(3)
    add_inline_runs(body, message)


def is_table_separator(line: str) -> bool:
    stripped = line.strip()
    return bool(re.match(r"^\|?[\s:-]+\|[\s|:-]*$", stripped))


def parse_table(lines: list[str], start: int) -> tuple[tuple[list[str], list[list[str]]], int] | tuple[None, int]:
    if start + 1 >= len(lines):
        return None, start
    header_line = lines[start].strip()
    separator_line = lines[start + 1].strip()
    if not (header_line.startswith("|") and header_line.endswith("|") and is_table_separator(separator_line)):
        return None, start

    def split_row(text: str) -> list[str]:
        return [cell.strip() for cell in text.strip().strip("|").split("|")]

    headers = split_row(header_line)
    rows: list[list[str]] = []
    index = start + 2
    while index < len(lines):
        stripped = lines[index].strip()
        if not (stripped.startswith("|") and stripped.endswith("|")):
            break
        rows.append(split_row(stripped))
        index += 1
    return (headers, rows), index


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, "17162C")
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.name = "Poppins SemiBold"
            run.font.size = Pt(9.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for row_index, row_values in enumerate(rows):
        row_cells = table.add_row().cells
        for col_index, value in enumerate(row_values):
            cell = row_cells[col_index]
            if row_index % 2 == 0:
                set_cell_shading(cell, "F8F9FC")
            paragraph = cell.paragraphs[0]
            add_inline_runs(paragraph, value)
            for run in paragraph.runs:
                run.font.size = Pt(9.5)


def build_status_callout(document: Document, items: list[str]) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F5F9")

    heading = cell.paragraphs[0]
    run = heading.add_run("Status project")
    run.font.name = "Poppins SemiBold"
    run.font.size = Pt(10.5)
    run.font.color.rgb = BRAND_NAVY

    for item in items:
        paragraph = cell.add_paragraph(style="List Bullet")
        add_inline_runs(paragraph, item)


def markdown_to_docx(input_markdown: Path, output_docx: Path) -> None:
    lines = input_markdown.read_text(encoding="utf-8").splitlines()
    metadata = extract_cover_metadata(lines, input_markdown)

    document = Document()
    configure_styles(document)
    build_cover(document, metadata)

    content_section = document.add_section(WD_SECTION_START.NEW_PAGE)
    content_section.top_margin = Cm(2.1)
    content_section.bottom_margin = Cm(1.9)
    content_section.left_margin = Cm(2.0)
    content_section.right_margin = Cm(2.0)

    build_summary_callout(document, metadata)
    document.add_paragraph("")

    index = 0
    chapter_number = 0
    first_title_skipped = False

    while index < len(lines):
        raw_line = lines[index]
        stripped = raw_line.strip()

        if not stripped:
            index += 1
            continue

        if stripped.startswith("# "):
            if not first_title_skipped:
                first_title_skipped = True
                index += 1
                continue

            chapter_number += 1
            label = document.add_paragraph(style="Subtitle")
            label.paragraph_format.space_before = Pt(12)
            label.paragraph_format.space_after = Pt(0)
            run = label.add_run(f"HOOFDSTUK {chapter_number:02d}")
            run.font.name = "Poppins Light"
            run.font.size = Pt(9)
            run.font.color.rgb = BRAND_MUTED

            heading = document.add_paragraph(style="Heading 1")
            add_inline_runs(heading, stripped[2:])
            index += 1
            continue

        if stripped.startswith("## "):
            chapter_number += 1
            label = document.add_paragraph(style="Subtitle")
            label.paragraph_format.space_before = Pt(12)
            label.paragraph_format.space_after = Pt(0)
            run = label.add_run(f"HOOFDSTUK {chapter_number:02d}")
            run.font.name = "Poppins Light"
            run.font.size = Pt(9)
            run.font.color.rgb = BRAND_MUTED

            heading = document.add_paragraph(style="Heading 1")
            add_inline_runs(heading, stripped[3:])
            index += 1
            continue

        if stripped.startswith("### "):
            heading = document.add_paragraph(style="Heading 2")
            add_inline_runs(heading, stripped[4:])
            index += 1
            continue

        if stripped.startswith("#### "):
            heading = document.add_paragraph(style="Heading 3")
            add_inline_runs(heading, stripped[5:])
            index += 1
            continue

        if re.match(r"^_{5,}$", stripped):
            add_separator(document)
            index += 1
            continue

        if stripped == "VS":
            add_vs_label(document)
            index += 1
            continue

        if stripped in CALLOUT_LABELS:
            add_special_label(document, stripped)
            index += 1
            continue

        if stripped == "STATUS PROJECT":
            items: list[str] = []
            index += 1
            while index < len(lines):
                candidate = lines[index].strip()
                if not candidate:
                    index += 1
                    continue
                if not candidate.startswith("- "):
                    break
                items.append(candidate[2:].strip())
                index += 1
            build_status_callout(document, items)
            continue

        table_data, next_index = parse_table(lines, index)
        if table_data is not None:
            headers, rows = table_data
            add_table(document, headers, rows)
            index = next_index
            continue

        if stripped.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline_runs(paragraph, stripped[2:])
            index += 1
            continue

        if re.match(r"^\d+\.\s+", stripped):
            paragraph = document.add_paragraph(style="List Number")
            add_inline_runs(paragraph, re.sub(r"^\d+\.\s+", "", stripped))
            index += 1
            continue

        if re.match(r"^\*\*(.+?):\*\*\s*(.+)$", stripped):
            key, value = re.match(r"^\*\*(.+?):\*\*\s*(.+)$", stripped).groups()
            paragraph = document.add_paragraph(style="Normal")
            add_inline_runs(paragraph, f"**{key}:** {value}")
            index += 1
            continue

        add_body_paragraph(document, stripped)
        index += 1

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_docx))


def merge_export_args(args: argparse.Namespace) -> argparse.Namespace:
    export_args = argparse.Namespace()
    export_args.input_file = build_output_docx_path(args.input_markdown, args.output_docx)
    export_args.folder_id = args.folder_id or os.getenv(
        "GOOGLE_DRIVE_FOLDER_ID",
        DEFAULT_GOOGLE_DRIVE_FOLDER_ID,
    )
    export_args.name = args.name
    export_args.service_account_file = args.service_account_file or os.getenv("GOOGLE_SERVICE_ACCOUNT_FILE")
    export_args.client_secrets_file = args.client_secrets_file or os.getenv("GOOGLE_OAUTH_CLIENT_SECRET_FILE")
    export_args.token_file = args.token_file or os.getenv("GOOGLE_OAUTH_TOKEN_FILE", "tmp/google-oauth-token.json")
    export_args.oauth_mode = args.oauth_mode or os.getenv("GOOGLE_OAUTH_MODE", "local-server")
    return export_args


def main() -> None:
    load_environment()
    args = parse_args()
    if not args.input_markdown.exists():
        raise SystemExit(f"Input markdown does not exist: {args.input_markdown}")

    export_args = merge_export_args(args)
    markdown_to_docx(args.input_markdown, export_args.input_file)
    export_args.name = export_args.name or args.input_markdown.stem
    validate_export_args(export_args)
    credentials = build_credentials(export_args)
    created_file = upload_google_doc(export_args, credentials)

    print(f"Generated DOCX: {export_args.input_file}")
    print(f"Created Google Doc: {created_file['name']}")
    print(f"File ID: {created_file['id']}")
    print(f"Open: {created_file['webViewLink']}")


if __name__ == "__main__":
    main()
