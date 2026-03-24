from __future__ import annotations

import argparse
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    KeepTogether,
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

try:
    from PIL import Image, ImageDraw, ImageFont
except ImportError:
    Image = None
    ImageDraw = None
    ImageFont = None


TITLE = "FavorCool SEO & GEO Analyse"
SUBTITLE = "2026-03-14"
BRAND_NAVY = "#17162C"
BRAND_NAVY_RGB = (0x17, 0x16, 0x2C)
BRAND_BLUEGREY = "#5C667F"
BRAND_TEXT = "#242938"
BRAND_MUTED = "#6B7388"
BRAND_SURFACE = "#F4F5F9"
BRAND_LINE = "#D8DCE6"
BRAND_WHITE = "#FFFFFF"
FONT_DIR = Path(__file__).resolve().parents[1] / "assets" / "fonts"

META = [
    "Bedrijf: FavorCool",
    "Website: https://www.favor-cool.be/",
    "Auditdatum: 14 maart 2026",
    "Focusdiensten: airco, warmtepompen, zonnepanelen",
    "Doelregio's: Antwerpen, Oost-Vlaanderen, West-Vlaanderen",
]

SEO_SCORE_ROWS = [
    ("Technische SEO", "6/10", "Degelijke basis, maar overlap in metadata en structured data remmen groei."),
    ("Content SEO", "6/10", "Sterke commerciële basis, maar nog weinig kenniscontent en clusterdiepte."),
    ("Local SEO", "4/10", "Vestigingen zijn aanwezig, maar regionale SEO-architectuur ontbreekt."),
    ("Mobile SEO", "5/10", "Mobiel bruikbaar, maar duidelijk te zwaar door scripts en CSS."),
    ("Structured Data", "5/10", "Aanwezig, maar te beperkt en deels onjuist ingevuld."),
    ("Trust Signals", "7/10", "Sterke commerciële geloofwaardigheid via reviews, claims en aanbodbreedte."),
]

COMP_ROWS = [
    ("FavorCool", "61", "1", "Geen echte provinciehubs", "Sterk commercieel, lokaal nog te zwak"),
    ("EnergyKing", "274", "134", "Ja", "Contentmachine met regionale schaal"),
    ("NextPower", "26", "1", "Beperkt", "Compacte, heldere servicestructuur"),
]

ROADMAP = {
    "Maand 1-2": [
        "metadata en paginarollen opschonen",
        "contactpagina corrigeren",
        "structured data verbeteren",
        "performance quick wins doorvoeren",
        "eerste lokale pagina's voor Antwerpen bouwen",
    ],
    "Maand 3-4": [
        "provinciehubs voor Oost- en West-Vlaanderen toevoegen",
        "eerste service + stad pagina's lanceren",
        "eerste 4 blogartikels publiceren",
        "FAQ-secties toevoegen aan kernpagina's",
    ],
    "Maand 5-6": [
        "contentcluster uitbreiden",
        "projectcases toevoegen",
        "lokale citations en backlinks opbouwen",
        "beste pagina's verder optimaliseren op basis van data",
    ],
}

TOP5 = [
    "Los de overlap op tussen airco-landingspagina's.",
    "Bouw lokale landingspagina's voor Antwerpen, Oost-Vlaanderen en West-Vlaanderen.",
    "Verbeter de mobiele performance door scripts en CSS te reduceren.",
    "Bouw een contentcluster rond prijs, vergelijking en premies.",
    "Versterk lokale trust via reviews, cases en betere structured data.",
]

IMPACT_ROWS = [
    ("Extra organische bezoekers", "+400 tot +700 / maand"),
    ("Extra leads", "6 tot 17 / maand"),
    ("Extra projecten", "1 tot 3 / maand"),
    ("Potentiële extra omzet", "EUR 6.000 tot EUR 18.000 / maand"),
]

CWV_ROWS = [
    ("LCP", "< 2,5 s", "17,6 s mobiel", "Te traag, vooral door front-end belasting"),
    ("INP", "< 200 ms", "geen velddata; TBT 150 ms", "Waakzaam, maar niet het grootste risico"),
    ("CLS", "< 0,10", "0,00 mobiel", "Stabiel"),
]

KEYWORD_ROWS = [
    ("airco installateur Antwerpen", "middel", "lokale dienstpagina"),
    ("warmtepomp installateur Antwerpen", "middel", "lokale dienstpagina"),
    ("zonnepanelen Antwerpen", "middel", "lokale dienstpagina"),
    ("airco prijs", "sterk", "prijscluster verfijnen"),
    ("airco zonder buitenunit", "sterk", "vergelijking en prijs combineren"),
    ("warmtepomp premies 2026", "zwak", "nieuw contentartikel bouwen"),
]

LOCAL_ROWS = [
    ("Antwerpen", "hoog", "eerste provinciehub en servicepagina's"),
    ("Gent", "hoog", "airco + warmtepomp + zonnepanelen"),
    ("Brugge", "middel", "zonnepanelen en warmtepomp"),
    ("Mechelen", "middel", "airco en warmtepomp"),
    ("Aalst", "middel", "warmtepomp en zonnepanelen"),
    ("Sint-Niklaas", "middel", "airco en zonnepanelen"),
]


@dataclass
class Section:
    title: str
    paragraphs: list[str]
    bullets: list[str] | None = None


SECTIONS: list[Section] = [
    Section(
        "Executive summary",
        [
            "FavorCool heeft vandaag al een veel sterkere commerciële website dan veel lokale installateurs. De site is duidelijk gebouwd om bezoekers richting offerte, contact of advies te sturen.",
            "De grootste remmers zijn te weinig lokale SEO-structuur, te weinig kenniscontent, overlap tussen airco-landingspagina's, zwakke mobiele performance en te beperkte structured data.",
            "De groeikans zit dus niet in helemaal opnieuw beginnen, maar in een slimme tweede fase: lokaler, scherper en inhoudelijk dieper werken.",
        ],
        [
            "sterke commerciële opbouw met duidelijke leadfocus",
            "nog te weinig regionale dekking voor Antwerpen, Oost- en West-Vlaanderen",
            "grote winst mogelijk via lokale pagina's, contentclusters en technische optimalisatie",
        ],
    ),
    Section(
        "Huidige situatie",
        [
            "FavorCool werkt vandaag als een marketingwebsite met leadgeneratie-elementen. Dat is sterker dan een brochure-site, omdat er duidelijke intentiepagina's en conversieflows aanwezig zijn.",
            "Sterke punten zijn de homepage, de dienstpagina's voor airco, warmtepompen en zonnepanelen, en de commerciële pagina's rond prijs en offerte.",
            "Wat ontbreekt, is vooral een structurele SEO-laag voor lokale zoekintentie en expertisecontent.",
        ],
        [
            "commerciële pagina's zoals airco prijs, airco kopen en airco offerte aanvragen",
            "product- en merkpagina's voor extra commerciële dekking",
            "blog aanwezig maar praktisch leeg",
            "geen provincie- of stadspagina's voor de gevraagde doelregio's",
        ],
    ),
    Section(
        "Grootste SEO-problemen",
        [
            "De belangrijkste SEO-problemen zitten niet in een gebrek aan basis, maar in de kwaliteit van de volgende stap.",
            "Meerdere airco-subpagina's lijken inhoudelijk te dicht op elkaar te zitten. Dat maakt het voor zoekmachines moeilijker om exact te begrijpen welke pagina voor welke intentie moet ranken.",
        ],
        [
            "geen lokale SEO-architectuur voor Antwerpen, Oost-Vlaanderen en West-Vlaanderen",
            "overlap tussen airco-landingspagina's",
            "nauwelijks blog- en kenniscontent",
            "mobiele performance duidelijk te zwak",
            "structured data te beperkt en deels onjuist",
        ],
    ),
    Section(
        "Grootste groeikansen",
        [
            "De snelste groeikansen liggen in lokale SEO en content. FavorCool heeft al een geloofwaardig merk en aanbod, maar moet dat nu beter vertalen naar lokale zoekintentie.",
            "Vooral pagina's rond airco, warmtepompen en zonnepanelen in combinatie met steden en provincies kunnen hier snel extra verkeer en aanvragen genereren.",
        ],
        [
            "lokale landingspagina's voor Antwerpen, Gent, Brugge, Mechelen, Aalst en Sint-Niklaas",
            "dienst + stad combinaties met hoge koopintentie",
            "content rond prijs, premies, vergelijking en keuzehulp",
            "answer-ready GEO-content voor AI-platformen",
        ],
    ),
    Section(
        "Benchmark met concurrenten",
        [
            "FavorCool heeft een sterkere commerciële basis dan veel kleine installateurs, maar wordt inhoudelijk overklast door spelers met meer contentvolume en meer regionale SEO-schaal.",
            "De vergelijking met EnergyKing en NextPower toont dat FavorCool vooral moet groeien in contentritme, lokale architectuur en thematische autoriteit.",
        ],
        [
            "EnergyKing scoort sterk op schaal en blogvolume",
            "NextPower heeft een compactere maar nette servicestructuur",
            "FavorCool heeft zelf al goede commerciële diepte, maar benut die nog onvoldoende lokaal",
        ],
    ),
    Section(
        "Technische quick wins",
        [
            "De technische quick wins zijn relatief duidelijk. De server reageert snel, dus het hoofdprobleem zit niet in hosting maar in front-end gewicht.",
            "Vooral overbodige scripts, third-party widgets en niet-kritische CSS wegen op de mobiele ervaring.",
        ],
        [
            "titles en meta descriptions per airco-pagina uniek maken",
            "contactpagina voorzien van een duidelijke H1",
            "structured data aanvullen met service, FAQ en correcte bedrijfsdata",
            "scripts en CSS reduceren voor betere mobiele laadtijd",
        ],
    ),
    Section(
        "Core Web Vitals in mensentaal",
        [
            "LCP toont hoe snel het grootste zichtbare element geladen is. Daar scoort FavorCool mobiel zwak op.",
            "INP toont hoe snel de site reageert op interactie. Dat lijkt technisch beheersbaar, maar de scriptbelasting is wel duidelijk te hoog.",
            "CLS meet visuele verschuivingen tijdens het laden. Dat onderdeel oogt relatief stabiel.",
        ],
        [
            "server is snel, dus hosting is niet het hoofdprobleem",
            "main thread werk is te zwaar op mobiel",
            "third-party scripts veroorzaken een groot deel van de vertraging",
        ],
    ),
    Section(
        "Lokale SEO-kansen",
        [
            "Lokale SEO is vandaag de grootste commerciële hefboom. FavorCool heeft al geloofwaardige vestigingen, maar die zijn nog niet strategisch uitgewerkt als zoekstructuur.",
            "Een duidelijke opbouw met provinciehubs, stadspagina's en dienst + stad combinaties kan hier het verschil maken.",
        ],
        [
            "Antwerpen als eerste prioriteit",
            "daarna Oost-Vlaanderen en West-Vlaanderen",
            "reviews, cases en servicegebieden sterker per regio koppelen",
        ],
    ),
    Section(
        "Google Business en trust",
        [
            "FavorCool heeft op de site al sterke trustelementen zoals reviewclaims, ervaringsclaims en meerdere productlijnen. De volgende stap is dat vertrouwen lokaler en zichtbaarder maken.",
            "Voor lokale SEO moeten reviews, projectfoto's en servicegebieden sterker gekoppeld worden aan concrete locaties en diensttypes.",
        ],
        [
            "reviews laten verwijzen naar dienst en regio",
            "foto-updates en projectposts per vestiging of regio",
            "trustsignalen opnemen op dienst- en landingspagina's",
        ],
    ),
    Section(
        "Keyword gap en prioritaire zoektermen",
        [
            "FavorCool heeft al een commerciële basis, maar laat nog belangrijke niet-merk zoekintentie liggen. Vooral zoektermen met lokale koopintentie en vraaggestuurde informatie zijn nog te weinig afgedekt.",
            "De combinatie van dienst + regio en prijs + keuzecontent is hier de snelste manier om extra relevant verkeer te winnen.",
        ],
        [
            "dienst + provincie",
            "dienst + stad",
            "prijs- en premievragen",
            "vergelijkingsvragen",
        ],
    ),
    Section(
        "Content- en GEO-aanpak",
        [
            "FavorCool moet werken rond drie kernpijlers: airco, warmtepompen en zonnepanelen. Rond elk van die pijlers hoort een cluster van prijsartikelen, vergelijkingen, FAQ's en keuzehulpen.",
            "Dat helpt tegelijk voor klassieke SEO en voor zichtbaarheid in AI-zoekmachines zoals ChatGPT, Gemini en Perplexity.",
        ],
        [
            "Wat kost een airco installatie in 2026?",
            "Monoblock airco vs split airco",
            "Wanneer kies je een warmtepomp?",
            "Zonnepanelen met thuisbatterij: loont dat nog?",
        ],
    ),
    Section(
        "Aanbevolen sitestructuur",
        [
            "De bestaande site heeft al een commerciële basis, maar mist een logische lokale SEO-laag. Die moet worden uitgebreid met provinciehubs, stadspagina's en service + stad combinaties.",
            "Zo ontstaat er een duidelijke interne linkstructuur van hoofdservice naar lokale pagina's naar content en offerteflow.",
        ],
        [
            "hoofdhubs voor Antwerpen, Oost-Vlaanderen en West-Vlaanderen",
            "stadspagina's voor prioritaire markten",
            "dienst + stad combinaties met hoge koopintentie",
            "blog- en FAQ-content die teruglinkt naar diensten",
        ],
    ),
    Section(
        "Aanbevolen contentplan",
        [
            "Om structureel topical authority op te bouwen, volstaat het niet om alleen commerciële pagina's te hebben. FavorCool heeft nood aan een ritme van nieuwe artikelen rond prijs, premies, vergelijking en gebruiksvragen.",
            "Twee sterke artikels per maand gedurende zes maanden volstaan al om een merkbaar verschil te maken in organisch bereik en AI-citeerbaarheid.",
        ],
        [
            "Wat kost een airco installatie in 2026?",
            "Monoblock airco vs split airco",
            "Wanneer kies je een warmtepomp?",
            "Premies voor warmtepompen in Vlaanderen",
            "Zonnepanelen met thuisbatterij",
            "Airco om te verwarmen: interessant of niet?",
        ],
    ),
    Section(
        "Roadmap en prioriteiten",
        [
            "De uitvoering hoeft niet in één keer. De grootste impact zit in een gefaseerde aanpak waarbij eerst de overlap en technische basis wordt aangepakt, en daarna de lokale en inhoudelijke schaal wordt opgebouwd.",
        ],
    ),
    Section(
        "Voorzichtige impactprognose",
        [
            "Met een conservatieve inschatting is een duidelijke uplift realistisch zodra de site beter scoort op niet-merk zoekintentie in de doelregio's.",
            "In deze markt hoeft SEO geen enorme aantallen leads op te leveren om commercieel rendabel te zijn. Eén tot drie extra projecten per maand maakt al een merkbaar verschil.",
        ],
    ),
    Section(
        "Conclusie",
        [
            "FavorCool heeft vandaag al een commerciële basis die sterker is dan die van veel lokale spelers. De site is dus geen zwakke basis, maar wel een basis die slimmer moet worden ingezet.",
            "De grootste hefboom zit in lokale SEO, contentdifferentiatie, kenniscontent en betere mobiele prestaties. Als die vier pijlers goed worden aangepakt, kan FavorCool uitgroeien tot een veel sterkere regionale speler in SEO en GEO.",
        ],
    ),
]


def clean_inline(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def register_pdf_fonts():
    font_map = {
        "regular": ("Poppins-Regular", FONT_DIR / "Poppins-Regular.ttf", "Helvetica"),
        "bold": ("Poppins-Bold", FONT_DIR / "Poppins-Bold.ttf", "Helvetica-Bold"),
        "light": ("Poppins-Light", FONT_DIR / "Poppins-Light.ttf", "Helvetica"),
        "semibold": ("Poppins-SemiBold", FONT_DIR / "Poppins-SemiBold.ttf", "Helvetica-Bold"),
    }
    resolved = {}
    for key, (font_name, font_path, fallback) in font_map.items():
        if font_path.exists():
            try:
                if font_name not in pdfmetrics.getRegisteredFontNames():
                    pdfmetrics.registerFont(TTFont(font_name, str(font_path)))
                resolved[key] = font_name
                continue
            except Exception:
                pass
        resolved[key] = fallback
    return resolved


def build_pdf_styles():
    fonts = register_pdf_fonts()
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="CoverEyebrow",
            parent=styles["BodyText"],
            fontName=fonts["light"],
            fontSize=10.5,
            leading=14,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#D6DAE4"),
            spaceAfter=10,
            uppercase=True,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CoverBrand",
            parent=styles["Title"],
            fontName=fonts["light"],
            fontSize=28,
            leading=32,
            alignment=TA_CENTER,
            textColor=colors.white,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CoverTitle",
            parent=styles["Title"],
            fontName=fonts["semibold"],
            fontSize=25,
            leading=30,
            alignment=TA_CENTER,
            textColor=colors.white,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CoverSub",
            parent=styles["BodyText"],
            fontName=fonts["regular"],
            fontSize=11.5,
            leading=14,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#CBD2E2"),
            spaceAfter=4,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Section",
            parent=styles["Heading1"],
            fontName=fonts["semibold"],
            fontSize=17.5,
            leading=23,
            textColor=colors.HexColor(BRAND_NAVY),
            spaceBefore=14,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Body",
            parent=styles["BodyText"],
            fontName=fonts["regular"],
            fontSize=11.2,
            leading=17,
            textColor=colors.HexColor(BRAND_TEXT),
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CalloutTitle",
            parent=styles["BodyText"],
            fontName=fonts["semibold"],
            fontSize=11.6,
            leading=14.5,
            textColor=colors.HexColor(BRAND_NAVY),
            spaceAfter=4,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Small",
            parent=styles["BodyText"],
            fontName=fonts["regular"],
            fontSize=9.3,
            leading=12.8,
            textColor=colors.HexColor(BRAND_MUTED),
        )
    )
    styles.add(
        ParagraphStyle(
            name="SectionLabel",
            parent=styles["BodyText"],
            fontName=fonts["light"],
            fontSize=9.5,
            leading=12,
            textColor=colors.HexColor(BRAND_BLUEGREY),
            spaceAfter=3,
            uppercase=True,
        )
    )
    return styles, fonts


def draw_cover_page(canvas, doc):
    canvas.saveState()
    canvas.setFillColor(colors.HexColor(BRAND_NAVY))
    canvas.rect(0, 0, doc.pagesize[0], doc.pagesize[1], fill=1, stroke=0)
    canvas.setFillColor(colors.HexColor("#5E667F"))
    canvas.setFont("Helvetica", 10)
    canvas.drawCentredString(doc.pagesize[0] / 2, 71 * mm, "favorcool")
    canvas.restoreState()


def draw_content_page(canvas, doc):
    canvas.saveState()
    canvas.setStrokeColor(colors.HexColor(BRAND_LINE))
    canvas.setLineWidth(0.6)
    canvas.line(18 * mm, doc.pagesize[1] - 15 * mm, doc.pagesize[0] - 18 * mm, doc.pagesize[1] - 15 * mm)
    canvas.setFillColor(colors.HexColor(BRAND_MUTED))
    canvas.setFont("Helvetica", 8.5)
    canvas.drawString(18 * mm, doc.pagesize[1] - 12 * mm, "FavorCool SEO & GEO Analyse")
    canvas.drawRightString(doc.pagesize[0] - 18 * mm, doc.pagesize[1] - 12 * mm, "2026-03-14")
    canvas.setFillColor(colors.HexColor(BRAND_MUTED))
    canvas.setFont("Helvetica", 8)
    canvas.drawRightString(doc.pagesize[0] - 18 * mm, 10 * mm, f"Pagina {doc.page}")
    canvas.restoreState()


def build_callout(title: str, body: str, styles):
    rows = [[Paragraph(clean_inline(title), styles["CalloutTitle"])]]
    if body:
        rows.append([Paragraph(clean_inline(body), styles["Body"])])
    tbl = Table(rows, colWidths=[174 * mm])
    tbl.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(BRAND_SURFACE)),
                ("LINEBEFORE", (0, 0), (0, -1), 2.4, colors.HexColor(BRAND_NAVY)),
                ("BOX", (0, 0), (-1, -1), 0.4, colors.HexColor(BRAND_LINE)),
                ("LEFTPADDING", (0, 0), (-1, -1), 12),
                ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ("TOPPADDING", (0, 0), (-1, -1), 9),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 9),
            ]
        )
    )
    return tbl


def build_three_col_score_table(styles):
    data = [
        [
            Paragraph("<b>Onderdeel</b>", styles["Small"]),
            Paragraph("<b>Score</b>", styles["Small"]),
            Paragraph("<b>Uitleg</b>", styles["Small"]),
        ]
    ]
    for name, score, desc in SEO_SCORE_ROWS:
        data.append(
            [
                Paragraph(clean_inline(name), styles["Small"]),
                Paragraph(clean_inline(score), styles["Small"]),
                Paragraph(clean_inline(desc), styles["Small"]),
            ]
        )
    tbl = Table(data, colWidths=[40 * mm, 18 * mm, 116 * mm], repeatRows=1)
    tbl.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(BRAND_NAVY)),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor(BRAND_LINE)),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor(BRAND_SURFACE)]),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    return tbl


def build_comp_table(styles):
    data = [
        [
            Paragraph("<b>Bedrijf</b>", styles["Small"]),
            Paragraph("<b>URL's</b>", styles["Small"]),
            Paragraph("<b>Blogposts</b>", styles["Small"]),
            Paragraph("<b>Regiohubs</b>", styles["Small"]),
            Paragraph("<b>Opmerking</b>", styles["Small"]),
        ]
    ]
    for row in COMP_ROWS:
        data.append([Paragraph(clean_inline(cell), styles["Small"]) for cell in row])
    tbl = Table(data, colWidths=[28 * mm, 18 * mm, 20 * mm, 28 * mm, 80 * mm], repeatRows=1)
    tbl.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(BRAND_NAVY)),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor(BRAND_LINE)),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor(BRAND_SURFACE)]),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    return tbl


def build_impact_table(styles):
    data = [
        [
            Paragraph("<b>Metric</b>", styles["Small"]),
            Paragraph("<b>Voorzichtige inschatting</b>", styles["Small"]),
        ]
    ]
    for row in IMPACT_ROWS:
        data.append([Paragraph(clean_inline(cell), styles["Small"]) for cell in row])
    tbl = Table(data, colWidths=[65 * mm, 109 * mm], repeatRows=1)
    tbl.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(BRAND_NAVY)),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor(BRAND_LINE)),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor(BRAND_SURFACE)]),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    return tbl


def build_four_col_table(styles, headers, rows, widths):
    data = [[Paragraph(f"<b>{clean_inline(h)}</b>", styles["Small"]) for h in headers]]
    for row in rows:
        data.append([Paragraph(clean_inline(cell), styles["Small"]) for cell in row])
    tbl = Table(data, colWidths=widths, repeatRows=1)
    tbl.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(BRAND_NAVY)),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor(BRAND_LINE)),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor(BRAND_SURFACE)]),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    return tbl


def pdf_bullets(items: Iterable[str], styles):
    return ListFlowable(
        [ListItem(Paragraph(clean_inline(item), styles["Body"])) for item in items],
        bulletType="bullet",
        leftIndent=14,
        bulletFontName="Helvetica",
        bulletFontSize=8,
    )


def build_pdf(output_path: Path):
    styles, _ = build_pdf_styles()
    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=22 * mm,
        bottomMargin=20 * mm,
        title=TITLE,
        author="Codex",
    )

    story = [
        Spacer(1, 60 * mm),
        Paragraph("SEO & GEO analyse 2026-03-14", styles["CoverEyebrow"]),
        Paragraph("FavorCool", styles["CoverBrand"]),
        Paragraph("SEO & GEO Analyse", styles["CoverTitle"]),
        Paragraph(SUBTITLE, styles["CoverSub"]),
        Spacer(1, 10 * mm),
    ]
    for line in META:
        story.append(Paragraph(clean_inline(line), styles["CoverSub"]))
    story.extend(
        [
            Spacer(1, 14 * mm),
            PageBreak(),
        ]
    )

    forced_breaks_after = {
        "Executive summary",
        "Huidige situatie",
        "Grootste SEO-problemen",
        "Benchmark met concurrenten",
        "Technische quick wins",
        "Benchmark met concurrenten",
        "Lokale SEO-kansen",
        "Keyword gap en prioritaire zoektermen",
        "Content- en GEO-aanpak",
        "Roadmap en prioriteiten",
    }

    for idx, section in enumerate(SECTIONS, start=1):
        story.append(Paragraph(f"hoofdstuk {idx:02d}", styles["SectionLabel"]))
        story.append(Paragraph(f"{idx}. {clean_inline(section.title)}", styles["Section"]))
        for paragraph in section.paragraphs:
            story.append(Paragraph(clean_inline(paragraph), styles["Body"]))
        if section.bullets:
            story.append(pdf_bullets(section.bullets, styles))
            story.append(Spacer(1, 4))

        if section.title == "Huidige situatie":
            story.append(
                build_callout(
                    "Snelle interpretatie",
                    "De site werkt vandaag als marketingwebsite met leadgeneratie-elementen, maar nog niet als volledig uitgewerkte regionale SEO-machine.",
                    styles,
                )
            )
        if section.title == "Grootste SEO-problemen":
            story.append(build_three_col_score_table(styles))
            story.append(Spacer(1, 10))
            story.append(
                build_callout(
                    "Wat dit commercieel betekent",
                    "Zonder lokale structuur en zonder duidelijke paginadifferentiatie blijft FavorCool vooral zichtbaar op merk en een beperkt aantal generieke termen, terwijl juist niet-merk zoekverkeer het groeipotentieel bevat.",
                    styles,
                )
            )
        if section.title == "Technische quick wins":
            story.append(Spacer(1, 8))
            story.append(
                build_four_col_table(
                    styles,
                    ["Metric", "Doel", "Huidig", "Interpretatie"],
                    CWV_ROWS,
                    [28 * mm, 24 * mm, 34 * mm, 88 * mm],
                )
            )
        if section.title == "Benchmark met concurrenten":
            story.append(build_comp_table(styles))
        if section.title == "Lokale SEO-kansen":
            story.append(Spacer(1, 8))
            story.append(
                build_four_col_table(
                    styles,
                    ["Locatie", "Prioriteit", "Aanpak"],
                    LOCAL_ROWS,
                    [36 * mm, 24 * mm, 114 * mm],
                )
            )
        if section.title == "Keyword gap en prioritaire zoektermen":
            story.append(Spacer(1, 8))
            story.append(
                build_four_col_table(
                    styles,
                    ["Zoekterm", "Potentieel", "Aanpak"],
                    KEYWORD_ROWS,
                    [74 * mm, 24 * mm, 76 * mm],
                )
            )
        if section.title == "Aanbevolen sitestructuur":
            story.append(
                build_callout(
                    "Doelbeeld",
                    "FavorCool groeit van een sterke commerciële site naar een lokale SEO-machine met provinciehubs, stadspagina's, service + stad combinaties en ondersteunende contentclusters.",
                    styles,
                )
            )
        if section.title == "Roadmap en prioriteiten":
            for phase, bullets in ROADMAP.items():
                story.append(
                    KeepTogether(
                        [
                            build_callout(phase, "", styles),
                            pdf_bullets(bullets, styles),
                            Spacer(1, 4),
                        ]
                    )
                )
            story.append(Paragraph("Top 5 acties met de grootste impact", styles["CalloutTitle"]))
            story.append(pdf_bullets(TOP5, styles))
        if section.title == "Voorzichtige impactprognose":
            story.append(build_impact_table(styles))
            story.append(
                build_callout(
                    "Belangrijke nuance",
                    "In deze sector hoeft SEO geen massaal volume aan extra leads op te leveren om een duidelijk commercieel rendement te hebben. Eén tot drie extra projecten per maand maakt al een merkbaar verschil.",
                    styles,
                )
            )
        if section.title in forced_breaks_after:
            story.append(PageBreak())

    doc.build(story, onFirstPage=draw_cover_page, onLaterPages=draw_content_page)


def set_cell_shading(cell, color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def set_doc_language(document: Document, lang: str = "nl-BE"):
    styles = document.styles
    for style_name in ["Normal", "Heading 1", "Heading 2", "Title", "Subtitle"]:
        if style_name in styles:
            style = styles[style_name]
            rpr = style.element.get_or_add_rPr()
            lang_el = OxmlElement("w:lang")
            lang_el.set(qn("w:val"), lang)
            rpr.append(lang_el)


def configure_docx_styles(document: Document):
    for style_name, font_name, size, color, bold in [
        ("Normal", "Poppins", Pt(10.8), RGBColor(0x24, 0x29, 0x38), False),
        ("Title", "Poppins Light", Pt(24), RGBColor(0x17, 0x16, 0x2C), False),
        ("Subtitle", "Poppins", Pt(12), RGBColor(0x6B, 0x73, 0x88), False),
        ("Heading 1", "Poppins SemiBold", Pt(18), RGBColor(0x17, 0x16, 0x2C), True),
        ("Heading 2", "Poppins SemiBold", Pt(14), RGBColor(0x17, 0x16, 0x2C), True),
    ]:
        style = document.styles[style_name]
        style.font.name = font_name
        style.font.size = size
        style.font.color.rgb = color
        style.font.bold = bold


def docx_heading(document: Document, text: str, level: int = 1):
    p = document.add_paragraph(style=f"Heading {level}")
    run = p.add_run(text)
    run.font.name = "Poppins SemiBold"
    run.font.color.rgb = RGBColor(*BRAND_NAVY_RGB)
    run.bold = True
    return p


def docx_bullets(document: Document, items: Iterable[str]):
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.name = "Poppins"
        run.font.size = Pt(10.5)


def docx_table(document: Document, headers: list[str], rows: list[tuple[str, ...]]):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        set_cell_shading(hdr_cells[i], "17162C")
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.bold = True
                run.font.name = "Poppins SemiBold"
                run.font.size = Pt(9.5)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            for p in cells[i].paragraphs:
                for run in p.runs:
                    run.font.name = "Poppins"
                    run.font.size = Pt(9.5)
    return table


def docx_callout(document: Document, title: str, body: str):
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F5F9")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = p.add_run(title)
    title_run.font.name = "Poppins SemiBold"
    title_run.font.size = Pt(10.5)
    title_run.font.color.rgb = RGBColor(*BRAND_NAVY_RGB)
    if body:
        p = cell.add_paragraph()
        run = p.add_run(body)
        run.font.name = "Poppins"
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0x24, 0x29, 0x38)
    return table


def build_cover_image(output_path: Path):
    if Image is None or ImageDraw is None or ImageFont is None:
        return False

    output_path.parent.mkdir(parents=True, exist_ok=True)
    image = Image.new("RGB", (1600, 2262), BRAND_NAVY)
    draw = ImageDraw.Draw(image)

    def load_font(name: str, size: int):
        font_path = FONT_DIR / name
        try:
            return ImageFont.truetype(str(font_path), size)
        except Exception:
            return ImageFont.load_default()

    eyebrow_font = load_font("Poppins-Light.ttf", 36)
    brand_font = load_font("Poppins-Light.ttf", 84)
    title_font = load_font("Poppins-SemiBold.ttf", 92)
    body_font = load_font("Poppins-Regular.ttf", 38)
    mark_font = load_font("Poppins-Light.ttf", 42)

    def centered_text(y: int, text: str, font, fill):
        bbox = draw.textbbox((0, 0), text, font=font)
        text_width = bbox[2] - bbox[0]
        draw.text(((1600 - text_width) / 2, y), text, font=font, fill=fill)

    centered_text(620, "SEO & GEO analyse 2026-03-14", eyebrow_font, "#D6DAE4")
    centered_text(760, "FavorCool", brand_font, "#FFFFFF")
    centered_text(890, "SEO & GEO Analyse", title_font, "#FFFFFF")
    centered_text(1015, "2026-03-14", body_font, "#CBD2E2")

    meta_lines = [
        "Bedrijf: FavorCool",
        "Website: https://www.favor-cool.be/",
        "Auditdatum: 14 maart 2026",
        "Focusdiensten: airco, warmtepompen, zonnepanelen",
        "Doelregio's: Antwerpen, Oost-Vlaanderen, West-Vlaanderen",
    ]
    for idx, line in enumerate(meta_lines):
        centered_text(1190 + (idx * 62), line, body_font, "#CBD2E2")

    centered_text(1950, "favorcool", mark_font, BRAND_BLUEGREY)
    image.save(output_path)
    return True


def set_table_width(table, width):
    table.autofit = False
    width_value = int(width)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(width_value // 635))
    for column in table.columns:
        column.width = width
    for row in table.rows:
        for cell in row.cells:
            cell.width = width


def build_docx_cover(document: Document, section):
    cover_table = document.add_table(rows=1, cols=1)
    cover_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(cover_table, section.page_width - section.left_margin - section.right_margin)
    row = cover_table.rows[0]
    row.height = Cm(22.7)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    cell = row.cells[0]
    set_cell_shading(cell, "17162C")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    first = cell.paragraphs[0]
    first.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = first.add_run("SEO & GEO ANALYSE 2026-03-14")
    run.font.name = "Poppins Light"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xD6, 0xDA, 0xE4)
    brand = cell.add_paragraph()
    brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = brand.add_run("FavorCool")
    run.font.name = "Poppins Light"
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    title = cell.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SEO & GEO Analyse")
    run.font.name = "Poppins SemiBold"
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    sub = cell.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("2026-03-14")
    run.font.name = "Poppins"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xCB, 0xD2, 0xE2)
    spacer = cell.add_paragraph()
    spacer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = spacer.add_run("")
    meta = cell.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("Antwerpen | Oost-Vlaanderen | West-Vlaanderen")
    run.font.name = "Poppins"
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0xCB, 0xD2, 0xE2)
    mark = cell.add_paragraph()
    mark.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = mark.add_run("favorcool")
    run.font.name = "Poppins Light"
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x5C, 0x66, 0x7F)


def build_docx(output_path: Path):
    document = Document()
    set_doc_language(document)
    configure_docx_styles(document)
    section = document.sections[0]
    section.top_margin = Cm(2.1)
    section.bottom_margin = Cm(1.9)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    cover_image_path = Path("tmp/docs/favorcool-docx-cover.png")
    if build_cover_image(cover_image_path):
        document.add_picture(str(cover_image_path), width=section.page_width - section.left_margin - section.right_margin)
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        build_docx_cover(document, section)

    document.add_page_break()
    docx_callout(
        document,
        "Belangrijkste boodschap",
        "FavorCool heeft al een sterke commerciële basis. De grootste winst zit nu in lokale SEO, contentdifferentiatie, betere mobiele performance en answer-ready GEO-content.",
    )
    document.add_paragraph()

    for idx, section_data in enumerate(SECTIONS, start=1):
        label = document.add_paragraph()
        label.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = label.add_run(f"HOOFDSTUK {idx:02d}")
        run.font.name = "Poppins Light"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x5C, 0x66, 0x7F)
        docx_heading(document, f"{idx}. {section_data.title}", 1)
        for paragraph in section_data.paragraphs:
            p = document.add_paragraph()
            run = p.add_run(paragraph)
            run.font.name = "Poppins"
            run.font.size = Pt(11)
        if section_data.bullets:
            docx_bullets(document, section_data.bullets)

        if section_data.title == "Grootste SEO-problemen":
            document.add_paragraph()
            docx_table(
                document,
                ["Onderdeel", "Score", "Uitleg"],
                SEO_SCORE_ROWS,
            )
            document.add_paragraph(
                "Wat dit commercieel betekent: zonder lokale structuur en zonder duidelijke paginadifferentiatie blijft FavorCool vooral zichtbaar op merk en een beperkt aantal generieke termen."
            )
        if section_data.title == "Technische quick wins":
            document.add_paragraph()
            docx_table(
                document,
                ["Metric", "Doel", "Huidig", "Interpretatie"],
                CWV_ROWS,
            )
        if section_data.title == "Benchmark met concurrenten":
            document.add_paragraph()
            docx_table(
                document,
                ["Bedrijf", "URL's", "Blogposts", "Regiohubs", "Opmerking"],
                COMP_ROWS,
            )
        if section_data.title == "Lokale SEO-kansen":
            document.add_paragraph()
            docx_table(
                document,
                ["Locatie", "Prioriteit", "Aanpak"],
                LOCAL_ROWS,
            )
        if section_data.title == "Keyword gap en prioritaire zoektermen":
            document.add_paragraph()
            docx_table(
                document,
                ["Zoekterm", "Potentieel", "Aanpak"],
                KEYWORD_ROWS,
            )
        if section_data.title == "Roadmap en prioriteiten":
            for phase, bullets in ROADMAP.items():
                docx_callout(document, phase, "")
                docx_bullets(document, bullets)
            docx_heading(document, "Top 5 acties met de grootste impact", 2)
            docx_bullets(document, TOP5)
        if section_data.title == "Voorzichtige impactprognose":
            document.add_paragraph()
            docx_table(
                document,
                ["Metric", "Voorzichtige inschatting"],
                IMPACT_ROWS,
            )
            document.add_paragraph(
                "Belangrijke nuance: in deze sector hoeft SEO geen massaal volume aan extra leads op te leveren om duidelijk commercieel rendement te hebben."
            )
        document.add_paragraph()

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("FavorCool SEO & GEO Analyse")
    run.font.name = "Poppins"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x6B, 0x73, 0x88)

    document.save(str(output_path))


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--pdf", type=Path, required=True)
    parser.add_argument("--docx", type=Path, required=True)
    args = parser.parse_args()
    args.pdf.parent.mkdir(parents=True, exist_ok=True)
    args.docx.parent.mkdir(parents=True, exist_ok=True)
    build_pdf(args.pdf)
    build_docx(args.docx)


if __name__ == "__main__":
    main()
